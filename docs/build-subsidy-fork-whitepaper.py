#!/usr/bin/env python3
"""Generate Bloodstone Subsidy Fork (1000 STONE) white paper."""

import importlib.util
import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/root/bloodstone-docs/Bloodstone-Subsidy-Fork-1000-White-Paper.docx"
COIN = 100_000_000
HALVING = 1_054_080
LEGACY_INITIAL = 100
NEW_INITIAL = 1000
FORK_HEIGHT = 12_000
INFLATION_FACTOR = 1.02956
INFLATION_BASE = 1_833_823_998


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D5E8F0")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def subsidy_at(height: int) -> float:
    if height < 1:
        return 0.0
    if height < FORK_HEIGHT:
        base = LEGACY_INITIAL
    else:
        base = NEW_INITIAL
    era = height // HALVING
    if era >= 64:
        return 0.0
    if era > 4:
        inflate = round(
            INFLATION_BASE
            * (math.pow(INFLATION_FACTOR, era - 3) - math.pow(INFLATION_FACTOR, era - 4))
        )
        n_subsidy = round(round(inflate / float(HALVING) * 100)) * (COIN // 100)
        n_subsidy = int(round(n_subsidy * base / 800))
        return n_subsidy / COIN
    return (base * COIN >> era) / COIN


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Bloodstone Subsidy Fork")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("1000 STONE Era-0 Schedule — Technical White Paper").font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\nVersion 1.0 · July 2026\n").font.size = Pt(11)
    meta.add_run("Activation block: 12,000 · bloodstonewallet.mytunnel.org").font.size = Pt(10)

    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Bloodstone mainnet launched in June 2026 with a 100 STONE proof-of-work block subsidy. "
        "In July 2026 the network will activate a scheduled consensus upgrade at block 12,000 that "
        "raises the era-0 subsidy to 1,000 STONE and applies standard halving from that base. "
        "All blocks below height 12,000 remain valid at 100 STONE; no chain reset or genesis change "
        "is required."
    )
    doc.add_paragraph(
        "This document specifies activation height, on-chain rules, the full subsidy table, pool "
        "alignment, operator upgrade steps, and economic implications relative to the original "
        "100 STONE schedule."
    )

    doc.add_heading("2. Why This Fork Exists", level=1)
    doc.add_paragraph(
        "The relaunch deliberately started at 100 STONE (one-eighth of legacy SpaceXpanse ROD’s "
        "800-coin era-0 reward) to signal a tighter issuance curve. After early mainnet operation, "
        "operators elected to increase miner incentives while preserving the long-run halving and "
        "inflation structure inherited from ROD consensus."
    )
    for item in [
        "Increase immediate PoW rewards from 100 to 1,000 STONE without invalidating mined history.",
        "Keep halving interval at 1,054,080 blocks (~2.7 years at ~80 s mean block time).",
        "Preserve five halving eras, then the long-run inflation phase (eras 5–63).",
        "Scale inflation-era subsidies proportionally (×1.25 vs legacy 800-ROD curve).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Consensus Rules", level=1)
    doc.add_heading("3.1 Activation", level=2)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Fork name", "Increased initial subsidy (internal: nIncreasedSubsidyHeight)"],
            ["Activation height", f"{FORK_HEIGHT:,}"],
            ["Pre-fork era-0 subsidy", f"{LEGACY_INITIAL} STONE"],
            ["Post-fork era-0 subsidy", f"{NEW_INITIAL} STONE"],
            ["Halving interval", f"{HALVING:,} blocks"],
            ["POST_ICO fork", "Active from block 1 (no 55,560-block bootstrap)"],
        ],
        [2.2, 3.3],
    )

    doc.add_heading("3.2 GetBlockSubsidy() behaviour", level=2)
    doc.add_paragraph(
        "For each block height n, the node selects a base initial subsidy:"
    )
    doc.add_paragraph(
        "• If n < 12,000: base = 100 STONE\n"
        "• If n ≥ 12,000: base = 1,000 STONE",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Halving index = floor(n / 1,054,080). For indices 0–4, subsidy = base >> index. "
        "For index ≥ 5, the SpaceXpanse inflation tranche formula applies, scaled by base / 800. "
        "ConnectBlock() rejects coinbases paying more than subsidy + fees (unchanged rule)."
    )

    doc.add_heading("4. Subsidy Schedule (Post-Fork)", level=1)
    doc.add_paragraph(
        "The table below uses the post-fork base (1,000 STONE) for era 0 onward. Blocks 1–11,999 "
        "historically paid 100 STONE in era 0."
    )

    era_rows = []
    for era in range(9):
        h = max(FORK_HEIGHT, era * HALVING + 1) if era == 0 else era * HALVING + 1
        if era == 0:
            h_label = f"1 – {HALVING - 1:,} (1,000 from {FORK_HEIGHT:,})"
        else:
            h_label = f"{era * HALVING:,} – {(era + 1) * HALVING - 1:,}"
        sub = subsidy_at(h)
        phase = "Halving" if era <= 4 else "Inflation"
        era_rows.append([str(era), h_label, f"{sub:g} STONE", phase])

    add_table(
        doc,
        ["Era", "Block height range", "Subsidy / block", "Phase"],
        era_rows,
        [0.5, 2.5, 1.2, 1.3],
    )

    doc.add_heading("4.1 Comparison: 100 vs 1,000 STONE base", level=2)
    compare_heights = [12000, 1054080, 2108160, 3162240, 4216320, 5270400]
    cmp_rows = []
    for h in compare_heights:
        era = h // HALVING
        old = (LEGACY_INITIAL * COIN >> era) / COIN if era <= 4 else subsidy_at(h)  # wrong for inflation with old base
        if era <= 4:
            old_sub = (LEGACY_INITIAL * COIN >> era) / COIN
        else:
            # old 100-base inflation
            inflate = round(
                INFLATION_BASE
                * (math.pow(INFLATION_FACTOR, era - 3) - math.pow(INFLATION_FACTOR, era - 4))
            )
            n = round(round(inflate / float(HALVING) * 100)) * (COIN // 100)
            old_sub = int(round(n * LEGACY_INITIAL / 800)) / COIN
        new_sub = subsidy_at(h)
        cmp_rows.append([str(h), str(era), f"{old_sub:g}", f"{new_sub:g}", f"{new_sub/old_sub:.1f}×" if old_sub else "—"])

    add_table(
        doc,
        ["Height", "Era", "100-base (hypothetical)", "1,000-base (active)", "Ratio"],
        cmp_rows,
        [1.1, 0.6, 1.2, 1.2, 0.6],
    )

    doc.add_heading("5. Supply Impact (Illustrative)", level=1)
    doc.add_paragraph(
        f"Between blocks {FORK_HEIGHT:,} and {HALVING - 1:,}, each block pays 1,000 STONE instead "
        f"of 100 — an extra 900 STONE per block for approximately {HALVING - FORK_HEIGHT:,} blocks "
        f"(~{((HALVING - FORK_HEIGHT) * 900 / 1e6):.2f} million STONE additional PoW issuance "
        "in era 0 alone, versus the counterfactual 100 STONE schedule)."
    )
    doc.add_paragraph(
        "Genesis premine (199,999,998 STONE) is unchanged. Halving dates are not reset; era 1 "
        "still begins at block 1,054,080 with 500 STONE under the new base."
    )

    doc.add_heading("6. Operator Upgrade Checklist", level=1)
    for item in [
        "Deploy Bloodstone Core build containing nIncreasedSubsidyHeight = 12000 and increasedInitialSubsidy = 1000 STONE.",
        "Restart bloodstoned on all seed nodes (64.188.22.190, 192.119.82.145) before block 12,000.",
        "Ensure miners produce coinbases ≤ new subsidy cap (pool reads getblockstats / formula).",
        "Set pool env: BLOODSTONE_INCREASED_SUBSIDY_HEIGHT=12000, BLOODSTONE_INCREASED_SUBSIDY_STONE=1000.",
        "Publish updated Qt/Android builds only if bundling node binaries; consensus change is daemon-side.",
        "Monitor debug.log past height 12,000 for 1,000 STONE coinbases via getblockstats.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Failure Modes", level=1)
    add_table(
        doc,
        ["Scenario", "Outcome", "Mitigation"],
        [
            ["All nodes upgraded before 12,000", "Smooth transition to 1,000 STONE", "Target state"],
            ["Minority hashrate on old binary at 12,000", "Chain stall until upgrade", "Upgrade seeds + pool first"],
            ["Pool pays 100 STONE formula after 12,000", "Underpayment to miners", "Deploy pool_block_subsidy.py"],
            ["Attempt config-only change", "No on-chain effect", "Must ship new bloodstoned"],
        ],
        [1.8, 1.8, 1.9],
    )

    doc.add_heading("8. Code References", level=1)
    for ref in [
        "chainparams.cpp — nIncreasedSubsidyHeight, increasedInitialSubsidy, initialSubsidy",
        "validation.cpp — GetEffectiveInitialSubsidy(), GetBlockSubsidy()",
        "pool_block_subsidy.py — effective_initial_stone(), schedule_preview()",
        "RPC: getblockstats <height> [\"subsidy\"] for live verification",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    doc.add_heading("9. Related Documents", level=1)
    for ref in [
        "Bloodstone Economic Model White Paper (July 2026)",
        "Bloodstone Subsidy Fork Release Notes (era-5 inflation scaling)",
        "Bloodstone Development Journey White Paper",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    fn = doc.add_paragraph()
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fn.add_run(
        "© 2026 Bloodstone · Consensus parameters are authoritative on-chain; "
        "calendar dates are estimates from ~80 s mean block time."
    ).font.size = Pt(9)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()