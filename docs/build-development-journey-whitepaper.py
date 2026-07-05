#!/usr/bin/env python3
"""Generate Bloodstone Development Journey white paper (.docx)."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "/root/bloodstone-docs/Bloodstone-Development-Journey-White-Paper.docx"


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D5E8F0")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    return p


def build():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bloodstone Development Journey")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "From the SpaceXpanse ROD Fork to the June 2026 Relaunch and Beyond"
    )
    sr.font.size = Pt(14)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\nWhite Paper · Version 1.0 · July 2026\n").font.size = Pt(11)
    meta.add_run("Bloodstone Core · bloodstonewallet.mytunnel.org").font.size = Pt(10)

    doc.add_page_break()

    # TOC placeholder
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Introduction",
        "3. Lineage: Bitcoin Core to SpaceXpanse ROD",
        "4. What ROD Inherited and Extended",
        "5. The Bloodstone Relaunch (June 2026)",
        "6. Genesis Block Specification",
        "7. Monetary Policy and the 0.7.x Core Releases",
        "8. Client and Platform Evolution",
        "9. Current Network State (July 2026)",
        "10. Lessons from the Field",
        "11. Roadmap Outlook",
        "12. References",
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1 Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Bloodstone is an independent proof-of-work blockchain that descends from the "
        "SpaceXpanse ROD core wallet — itself a long-running fork of Bitcoin Core enriched "
        "with multi-algorithm mining, on-chain naming, game integration, and auxiliary proof-of-work. "
        "In June 2026 the project executed a deliberate chain relaunch under the Bloodstone brand, "
        "resetting history at a new genesis block while preserving the technical architecture that "
        "made ROD suitable as a decentralised application platform."
    )
    doc.add_paragraph(
        "This document traces that development journey: the ROD heritage, the relaunch genesis "
        "and premine design, the Core 0.7.0 monetary-policy alignment, the ecosystem that grew "
        "around the chain (pool, Android node, Chain Mesh, web wallet), and the engineering "
        "incidents — including the v0.7.0 genesis-hex defect resolved in v0.7.1 — that shaped "
        "today's production network."
    )

    add_table(
        doc,
        ["Milestone", "Date / Era", "Significance"],
        [
            ["ROD mainnet (SpaceXpanse)", "Pre-2026", "800 ROD subsidy, legacy premine multisig, years of chain history"],
            ["Bloodstone relaunch genesis", "22 Jun 2026", "Height 0; independent STONE chain; 100 STONE era-0 reward"],
            ["Core 0.7.0 release", "Jul 2026", "Subsidy fork documentation; Windows Qt wallet; embedded genesis blob"],
            ["Core 0.7.1 hotfix", "Jul 2026", "Corrected genesis transaction serialization; fresh-install fix"],
            ["Mainnet growth", "Jul 2026", "~8,600+ blocks; pool, mesh, Android full-node path live"],
        ],
        [2.0, 1.2, 2.3],
    )

    # 2 Introduction
    doc.add_heading("2. Introduction", level=1)
    doc.add_paragraph(
        "Blockchain projects rarely begin on a blank slate. Bloodstone's story starts with ROD — "
        "the native coin of SpaceXpanse, marketed as the foundational layer of a multiverse platform "
        "for games, names, and decentralised applications. The rod-core-wallet repository (publicly "
        "associated with SpaceXpanse on GitHub) carried forward Bitcoin Core's UTXO model, peer networking, "
        "and wallet infrastructure, then layered consensus rules for NeoScrypt, SHA-256d merge-mining, "
        "and eventually Yespower, alongside name operations and game-aware RPC."
    )
    doc.add_paragraph(
        "By 2026 the Bloodstone team chose independence: a new genesis, a new ticker (STONE), a reduced "
        "but deliberate initial block subsidy, and a clean break from pre-relaunch chain data. The relaunch "
        "was not a mere rebrand — it was a reset of social contract, economics, and operational ownership, "
        "while reusing battle-tested C++ node code."
    )
    doc.add_paragraph(
        "Readers of companion documents — the Economic Model, Decentralized Network, Chain Mesh Storage, "
        "Live-Patchable Node, and Time Capsule white papers — will find here the chronological backbone "
        "that those specialised treatises assume."
    )

    # 3 Lineage
    doc.add_heading("3. Lineage: Bitcoin Core to SpaceXpanse ROD", level=1)
    doc.add_paragraph(
        "ROD core wallet follows the architectural patterns of Bitcoin Core: a full-node daemon "
        "(historically spacexpansed, now bloodstoned), a Qt GUI, CLI RPC tools, LevelDB-backed block "
        "and chainstate stores, and a script-based UTXO ledger. SpaceXpanse modifications introduced:"
    )
    for item in [
        "Auxiliary proof-of-work (auxpow) headers binding merge-mined work to parent chains.",
        "Dual- and triple-algorithm block production (NeoScrypt, SHA-256d, Yespower) with fork heights governing activation.",
        "On-chain name records — human-readable identifiers that can carry value and participate in game logic.",
        "Game RPC and ZMQ interfaces for decentralised applications to observe chain state in near real time.",
        "A 30-second target block interval (versus Bitcoin's 10 minutes), with halving every ~1,054,080 blocks.",
        "Legacy initial subsidy of 800 coins per block (ROD), with five halving eras then long-run inflation.",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph(
        "The README of the upstream rod-core-wallet repository (SpaceXpanse/rod-core-wallet) summarises ROD's "
        "purpose: secure value storage, transaction processing, mining, NFT/name support, and Play-to-Earn "
        "experiments. Bloodstone retains this codebase structure; binaries were renamed (bloodstoned, "
        "bloodstone-qt, bloodstone-cli) while internal build targets may still reference spacexpanse_* paths "
        "during migration."
    )

    # 4 What ROD Inherited
    doc.add_heading("4. What ROD Inherited and Extended", level=1)

    doc.add_heading("4.1 Consensus and Mining", level=2)
    doc.add_paragraph(
        "ROD's security model relies on proof-of-work across multiple algorithms. NeoScrypt targets CPU/GPU "
        "miners; SHA-256d enables merge-mining compatibility; Yespower extends CPU-friendly work at a defined "
        "fork height. Block headers embed a pure header plus algorithm-specific PoW data — a pattern Bloodstone "
        "preserved at relaunch."
    )

    doc.add_heading("4.2 Naming and Applications", level=2)
    doc.add_paragraph(
        "Unlike Bitcoin, ROD names are first-class ledger objects. Games and marketplaces use atomic name updates "
        "to trade in-game assets. Bloodstone mainnet ships with name history optional (off by default at relaunch) "
        "but retains the opcode and indexing hooks."
    )

    doc.add_heading("4.3 Legacy Economics (Pre-Relaunch)", level=2)
    add_table(
        doc,
        ["Parameter", "ROD (legacy)", "Notes"],
        [
            ["Initial PoW subsidy", "800 ROD", "Era-0 block reward before halvings"],
            ["POST_ICO bootstrap", "55,560 blocks × 1 coin", "Pre-ICO phase on legacy chain"],
            ["Halving interval", "1,054,080 blocks", "Carried forward to Bloodstone"],
            ["Inflation era", "Eras 5–63 (~3% per era)", "Unscaled formula; see §7 for Bloodstone adjustment"],
            ["Mainnet premine", "2-of-4 multisig (P2SH)", "SpaceXpanse founding-team custody"],
        ],
        [2.0, 1.5, 1.5],
    )

    # 5 Relaunch
    doc.add_heading("5. The Bloodstone Relaunch (June 2026)", level=1)
    doc.add_paragraph(
        "On 22 June 2026 Bloodstone activated an independent mainnet with a new genesis block. The coinbase "
        "message reads: \"22/Jun/2026: Bloodstone independent chain relaunch\". This event:"
    )
    for item in [
        "Discarded all pre-relaunch block files and chainstate (operators must wipe blocks/ and chainstate/).",
        "Introduced STONE as the user-facing unit (1 STONE = 1×10⁸ base units, same satoshi-style granularity as ROD).",
        "Set era-0 subsidy to 100 STONE — an order of magnitude below legacy 800 ROD, signalling a tighter issuance curve.",
        "Moved the premine to a single P2PKH address (SZNtmBMyx2Cr9VMrj5vk5EYTUn1naedu5N) instead of the legacy 2-of-4 multisig.",
        "Assigned new network magic bytes (0xa2, 0xf2, 0xf6, 0x93), P2P port 17333, RPC port 18332, and Bech32 HRP stone.",
        "Started POST_ICO rules at block 1 (no 55,560-block bootstrap on the new chain).",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph(
        "Qt and daemon builds write a relaunch marker file (.bloodstone_relaunch_genesis) after successful "
        "genesis connection, and offer to wipe stale SpaceXpanse-era folders when detected — reducing support "
        "burden for desktop users migrating datadir paths."
    )

    # 6 Genesis
    doc.add_heading("6. Genesis Block Specification", level=1)
    doc.add_paragraph(
        "The relaunch genesis is not recomputed at runtime from timestamp and nonce alone — doing so risks "
        "cross-platform floating-point or serialization drift. Instead, mainnet loads a canonical hex blob "
        "(LoadMainnetGenesisBlock) embedded in chainparams.cpp, with hard assertions on block hash and Merkle root."
    )

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Block hash", "df04225074039e630dad825b24818a695462bd19cd585131a0568f50e9bf71d0"],
            ["Merkle root", "3fed595c7afce69566ab84ed64a210719af8a0e20231a2666f282ecfb40b8ecb"],
            ["Timestamp (unix)", "1780569600"],
            ["Coinbase text", "22/Jun/2026: Bloodstone independent chain relaunch"],
            ["Premine output", "199,999,998 STONE"],
            ["Premine address", "SZNtmBMyx2Cr9VMrj5vk5EYTUn1naedu5N (P2PKH)"],
            ["PoW algorithm (era 0)", "NeoScrypt (nNonce 676154 in live blob)"],
        ],
        [2.2, 3.3],
    )

    doc.add_heading("6.1 The v0.7.0 Genesis Hex Defect", level=2)
    doc.add_paragraph(
        "Bloodstone Core 0.7.0 introduced the embedded genesis blob to eliminate reconstruction drift. "
        "However, the initial hex string omitted one zero byte in the coinbase input's null prevout hash (31 bytes "
        "instead of 32). Symptom on every fresh install:"
    )
    p = doc.add_paragraph()
    p.add_run("bad-txnmrklroot, hashMerkleRoot mismatch").italic = True

    doc.add_paragraph(
        "Nodes with existing synced data never re-validated genesis and continued operating. New Windows Qt "
        "installs failed immediately with \"A fatal internal error occurred\". Core 0.7.1 corrected the byte, "
        "added assert(BlockMerkleRoot(genesis) == genesis.hashMerkleRoot), and required users to reset blocks/ "
        "and chainstate/ after a failed 0.7.0 attempt. This incident underscores why genesis constants demand "
        "round-trip Merkle verification, not hash-only asserts."
    )

    # 7 Monetary
    doc.add_heading("7. Monetary Policy and the 0.7.x Core Releases", level=1)
    doc.add_paragraph(
        "Core 0.7.0 codified Bloodstone's issuance schedule for decades ahead. The critical divergence from "
        "legacy ROD is inflation scaling: after era 4, the SpaceXpanse formula would inflate subsidies as if "
        "the base were still 800 coins. Bloodstone applies a scale factor of 0.125 (100/800) so era 5 pays "
        "~6.62 STONE instead of ~52.95 STONE."
    )

    add_table(
        doc,
        ["Era", "Start block", "Subsidy / block", "Phase"],
        [
            ["0", "1", "100 STONE", "Halving era"],
            ["1", "1,054,080", "50 STONE", "Halving"],
            ["2", "2,108,160", "25 STONE", "Halving"],
            ["3", "3,162,240", "12.5 STONE", "Halving"],
            ["4", "4,216,320", "6.25 STONE", "Halving"],
            ["5", "5,270,400", "~6.62 STONE", "Scaled inflation"],
            ["64+", "—", "0", "PoW issuance ends"],
        ],
        [0.8, 1.3, 1.3, 1.1],
    )

    doc.add_paragraph(
        "The unified mining pool reads live subsidies via getblockstats; dashboard projections use the same "
        "parameters. A network-wide node upgrade is recommended before block 5,270,400 so on-chain consensus "
        "matches operator documentation."
    )

    # 8 Client evolution
    doc.add_heading("8. Client and Platform Evolution", level=1)

    doc.add_heading("8.1 Core Node and Desktop Wallet", level=2)
    for item in [
        "bloodstoned — full node daemon (systemd-managed on VPS seeders).",
        "bloodstone-qt v0.7.1 — Windows x64 static Qt 5.12 GUI; cross-compiled from Linux with mingw depends.",
        "bloodstone-cli / bloodstone-wallet — RPC and wallet utilities.",
        "Chain-reset PowerShell scripts published on the downloads CDN for datadir and blocks/chainstate recovery.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8.2 Mining Stack", level=2)
    doc.add_paragraph(
        "Pool workers on 192.119.82.145 run stratum bridges for NeoScrypt and Yespower, remote cpuminer fleets, "
        "and SV2-compatible backends. Android miner APK 1.3.23 restored local full-node mode by lowering storage "
        "gates, starting the foreground node service, and removing silent full→mesh downgrade in JavaScript."
    )

    doc.add_heading("8.3 Chain Mesh and Storage", level=2)
    doc.add_paragraph(
        "Chain Mesh provides chunked, searchable off-chain storage with glob-aware indexing — enabling large "
        "artefacts (APK mirrors, white papers, patch bundles) to propagate across mesh nodes. Search fixes in "
        "July 2026 corrected SQL LIKE escaping so patterns such as *Stuff_IN_HERE*.* return correct results."
    )

    doc.add_heading("8.4 Web Wallet and Identity", level=2)
    doc.add_paragraph(
        "bloodstonewallet.mytunnel.org hosts registration, dashboard, and downloads. July 2026 added X (Twitter) "
        "OAuth login (/wallet/x/login, callback, connect/disconnect), privacy and terms pages, and nginx "
        "redirects for legacy /privacy and /terms paths."
    )

    doc.add_heading("8.5 Documentation Suite", level=2)
    doc.add_paragraph(
        "Parallel white papers document economics, decentralised topology, mesh storage, live-patchable nodes, "
        "and time-capsule archival. This Development Journey paper ties them to a single timeline."
    )

    # 9 Current state
    doc.add_heading("9. Current Network State (July 2026)", level=1)
    doc.add_paragraph(
        "As of early July 2026 the public seed node reports mainnet height above 8,600 blocks, initial block "
        "download complete on infrastructure nodes, and era-0 subsidies of 100 STONE per block. Verification "
        "progress on a fully synced archival node remains low in percentage terms because the assumed chain "
        "length constant still reflects legacy ROD expectations — a cosmetic RPC artefact, not a sync failure."
    )

    add_table(
        doc,
        ["Parameter", "Live value"],
        [
            ["Network", "main"],
            ["Genesis hash", "df042250…71d0"],
            ["P2P port", "17333"],
            ["RPC port", "18332"],
            ["Bech32 HRP", "stone"],
            ["Address prefix (P2PKH)", "S (pubkey hash version 63)"],
            ["Recommended desktop build", "bloodstone-qt 0.7.1 (Windows)"],
            ["Recommended Android build", "bloodstone-miner-android 1.3.23"],
        ],
        [2.5, 3.0],
    )

    # 10 Lessons
    doc.add_heading("10. Lessons from the Field", level=1)
    lessons = [
        ("Embed canonical genesis bytes", "Recomputing genesis across platforms is fragile; verify Merkle roots at compile time."),
        ("Relaunch markers and UX", "Desktop wallets must detect pre-relaunch datadirs and offer safe wipes that preserve wallets/."),
        ("Publish pipeline completeness", "Binaries without index.html deployment left users on broken 0.7.0 links — release automation must ship HTML and checksums together."),
        ("Economics must be explicit", "Inherited inflation code paths need scaling when initial subsidy changes by an order of magnitude."),
        ("Fresh-install testing", "Long-running seeders mask genesis bugs; CI should spawn empty datadir smoke tests per release."),
    ]
    for title, body in lessons:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(body)

    # 11 Roadmap
    doc.add_heading("11. Roadmap Outlook", level=1)
    for item in [
        "Complete CLIENT_VERSION bump and reproducible Windows/Linux release parity for 0.7.1+.",
        "AssumeUTXO snapshots and assumevalid updates tuned to Bloodstone chain length (remove legacy ROD placeholders).",
        "HD wallet prefix refresh (BIP32 versions still marked FIXME in chainparams).",
        "Time Capsule + pruned fleet rollout per Time Capsule white paper.",
        "Live-patchable node bundles for zero-downtime bloodstoned upgrades.",
        "Further Android mesh integration and STONE wallet features (X-linked accounts, referral economy).",
    ]:
        add_bullet(doc, item)

    # 12 References
    doc.add_heading("12. References", level=1)
    refs = [
        "SpaceXpanse rod-core-wallet — https://github.com/SpaceXpanse/rod-core-wallet",
        "Bloodstone downloads — https://bloodstonewallet.mytunnel.org/downloads/",
        "Bloodstone Economic Model White Paper (July 2026)",
        "Bloodstone Decentralized Network White Paper (July 2026)",
        "Bloodstone Chain Mesh Storage White Paper v1.1 (July 2026)",
        "Bloodstone Subsidy Fork Release Notes — Core 0.7.0 (July 2026)",
        "chainparams.cpp — LoadMainnetGenesisBlock, mainnet consensus constants",
        "init.cpp — relaunch genesis marker (.bloodstone_relaunch_genesis)",
    ]
    for r in refs:
        add_bullet(doc, r)

    # Footer note
    doc.add_paragraph()
    fn = doc.add_paragraph()
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fn.add_run(
        "© 2026 Bloodstone · This document is informational and does not constitute investment advice."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()